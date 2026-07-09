#!/usr/bin/env python3
"""Convert a Hebrew (RTL) markdown file into a matching .docx file.

Usage:
    python3 md_to_docx.py <input.md> <output.docx>

Supports: # / ## / ### headings, **bold**, `code`, bullet/numbered lists,
> blockquotes, and pipe tables (rendered as real Word tables). All text is
rendered right-to-left with right-aligned paragraphs and RTL table layout.
"""
import re
import sys

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEBREW_FONT = "Arial"


def set_rtl_paragraph(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    run.font.name = HEBREW_FONT
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), HEBREW_FONT)


def set_table_rtl(table):
    tblPr = table._tbl.tblPr
    tblPr.append(OxmlElement('w:bidiVisual'))


def add_bold_runs(paragraph, text, force_bold=False, size=None):
    tokens = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = 'Courier New'
        else:
            run = paragraph.add_run(tok)
        if force_bold:
            run.bold = True
        if size:
            run.font.size = Pt(size)
        set_rtl_run(run)


def add_rtl_paragraph(doc, text, style=None, size=None):
    p = doc.add_paragraph(style=style)
    set_rtl_paragraph(p)
    add_bold_runs(p, text, size=size)
    return p


def parse_table_block(lines, start_idx):
    header = [c.strip() for c in lines[start_idx].strip().strip('|').split('|')]
    rows = [header]
    i = start_idx + 2
    while i < len(lines) and lines[i].strip().startswith('|'):
        rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
        i += 1
    return rows, i


def add_table(doc, rows):
    ncols = len(rows[0])
    table = doc.add_table(rows=0, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_rtl(table)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, cell_text in enumerate(row[:ncols]):
            cell = cells[c_idx]
            cell.paragraphs[0].text = ''
            p = cell.paragraphs[0]
            set_rtl_paragraph(p)
            add_bold_runs(p, cell_text, force_bold=(r_idx == 0), size=10)
    return table


def convert(src_path, out_path):
    with open(src_path, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    i, n = 0, len(lines)
    while i < n:
        stripped = lines[i].strip()

        if stripped == '':
            i += 1
            continue

        if stripped.startswith('# '):
            p = doc.add_heading(level=0)
            set_rtl_paragraph(p)
            add_bold_runs(p, stripped[2:].strip(), size=20)
            i += 1
            continue

        if stripped.startswith('### '):
            p = doc.add_heading(level=2)
            set_rtl_paragraph(p)
            add_bold_runs(p, stripped[4:].strip(), size=13)
            i += 1
            continue

        if stripped.startswith('## '):
            p = doc.add_heading(level=1)
            set_rtl_paragraph(p)
            add_bold_runs(p, stripped[3:].strip(), size=15)
            i += 1
            continue

        if stripped.startswith('> '):
            p = add_rtl_paragraph(doc, stripped[2:].strip())
            p.paragraph_format.left_indent = Cm(0.8)
            i += 1
            continue

        if stripped.startswith('|'):
            if i + 1 < n and re.match(r'^\|[\s:\-|]+\|?$', lines[i + 1].strip()):
                rows, next_i = parse_table_block(lines, i)
                add_table(doc, rows)
                doc.add_paragraph()
                i = next_i
                continue

        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            set_rtl_paragraph(p)
            add_bold_runs(p, stripped[2:].strip())
            i += 1
            continue

        if re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            set_rtl_paragraph(p)
            add_bold_runs(p, re.sub(r'^\d+\.\s', '', stripped))
            i += 1
            continue

        add_rtl_paragraph(doc, stripped)
        i += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print(__doc__)
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
